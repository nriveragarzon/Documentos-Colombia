# Importar funciones de los demás archivos
# Conexión 
# from conexion import obtener_conexion_snowflake
# Datos de exportaciones
# import datos_exportaciones as exportaciones

# Liberias 
#import os
#import json
#import snowflake.connector
#from snowflake.snowpark import Session
#import pandas as pd
#import numpy as np
#import time
#import re
#import warnings
#from docx import Document
#from docx.shared import Pt, RGBColor, Inches
#from docx.enum.style import WD_STYLE_TYPE
#from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
#from docx.oxml import parse_xml, OxmlElement
#from docx.oxml.ns import nsdecls, qn

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

################################################
# FUNCIONES PARA DEFINIR ESTILOS Y CREAR OBJETOS
################################################

# Función auxiliar para personalizar estilos
def customize_style(style, font_name, font_size, font_color, bold=False):
    style.font.name = font_name
    style.font.size = font_size
    style.font.color.rgb = font_color
    style.font.bold = bold
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Función para definir los estilos del documentos
def estilos(doc: Document):
    """
    Define y aplica estilos personalizados al documento proporcionado.
    
    Args:
    doc (Document): El documento al que se añadirán los estilos.
    """
  
    # Personalizar el estilo 'Title'
    title_style = doc.styles['Title']
    customize_style(title_style, 'Century Gothic', Pt(16), RGBColor(0, 32, 96), bold=True)

    # Personalizar el estilo 'Heading 1'
    heading1_style = doc.styles['Heading 1']
    customize_style(heading1_style, 'Century Gothic', Pt(14), RGBColor(0, 32, 96), bold=True)

    # Personalizar el estilo 'Heading 2'
    heading2_style = doc.styles['Heading 2']
    customize_style(heading2_style, 'Century Gothic', Pt(12), RGBColor(0, 32, 96), bold=True)

    # Personalizar el estilo 'Heading 3'
    heading3_style = doc.styles['Heading 3']
    customize_style(heading3_style, 'Century Gothic', Pt(12), RGBColor(0, 32, 96), bold=True)

    # Personalizar el estilo 'Normal'
    normal_style = doc.styles['Normal']
    customize_style(normal_style, 'Century Gothic', Pt(11), RGBColor(0, 0, 0))

    # Personalizar el estilo 'Table'
    table_style = doc.styles['Table Grid']
    table_font = table_style.font
    table_font.name = 'Century Gothic'

# Función para agregar un encabezado
def add_heading(doc: Document, text: str, level: int, style: str):
    """
    Agrega un encabezado al documento.

    Args:
    doc (Document): El documento al que se añadirá el encabezado.
    text (str): El texto del encabezado.
    level (int): El nivel del encabezado.
    style (str): El estilo del encabezado.
    """
    doc.add_heading(text, level=level).style = doc.styles[style]

# Función para agregar un párrafo
def add_paragraph(doc: Document, text: str, style: str):
    """
    Agrega un párrafo al documento.

    Args:
    doc (Document): El documento al que se añadirá el párrafo.
    text (str): El texto del párrafo.
    style (str): El estilo del párrafo.
    """
    p = doc.add_paragraph(text)
    p.style = doc.styles[style]


# Función para bordes de tablas
def set_cell_border(cell, **kwargs):
    """
    Establece los bordes de una celda.
    
    Uso:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "#00FF00", "space": "0"},
        left={"sz": 24, "val": "single", "color": "#0000FF", "space": "0"},
        right={"sz": 24, "val": "single", "color": "#000000", "space": "0"},
    )

    Args:
    cell: La celda a la que se aplicarán los bordes.
    kwargs: Un diccionario con las especificaciones para los bordes (tamaño, valor, color y espacio).
    """
    tcPr = cell._element.get_or_add_tcPr()

    # Verificar si ya existen bordes en la celda
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # Aplicar los bordes según las especificaciones proporcionadas
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))

# Función para crear tablas
def add_table(doc: Document, dataframe: pd.DataFrame, style: str):
    """
    Agrega una tabla al documento a partir de un DataFrame.

    Args:
    doc (Document): El documento al que se añadirá la tabla.
    dataframe (DataFrame): El DataFrame que se convertirá en tabla.
    style (str): El estilo de la tabla.
    """
    if isinstance(dataframe, pd.DataFrame):
        table = doc.add_table(rows=1, cols=len(dataframe.columns))
        table.style = doc.styles[style]
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True  # Desactivar ajuste automático para controlar el ancho de las celdas

        hdr_cells = table.rows[0].cells
        for i, column in enumerate(dataframe.columns):
            hdr_cells[i].text = str(column)
            hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hdr_cells[i].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), "#215E99")
            hdr_cells[i]._element.get_or_add_tcPr().append(shading_elm)
            set_cell_border(
                hdr_cells[i],
                top={"sz": 1, "val": "single", "color": "000000"},
                bottom={"sz": 1, "val": "single", "color": "000000"},
                left={"sz": 1, "val": "single", "color": "000000"},
                right={"sz": 1, "val": "single", "color": "000000"},
            )
        
        for index, row in dataframe.iterrows():
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)
                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row_cells[i].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
                if index == len(dataframe) - 1:  # Última fila
                    row_cells[i].paragraphs[0].runs[0].bold = True
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "#DAE9F7")
                    row_cells[i]._element.get_or_add_tcPr().append(shading_elm)
                    row_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                set_cell_border(
                    row_cells[i],
                    top={"sz": 1, "val": "single", "color": "000000"},
                    bottom={"sz": 1, "val": "single", "color": "000000"},
                    left={"sz": 1, "val": "single", "color": "000000"},
                    right={"sz": 1, "val": "single", "color": "000000"},
                )

        # Ajustar el ancho de las celdas para que la tabla ocupe todo el ancho de las márgenes del documento
        for row in table.rows:
            for cell in row.cells:
                cell.width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

    else:
        print(f"El valor proporcionado no es un DataFrame: {dataframe}")

# Función para agregar tabla de contenidos
def agregar_tabla_contenidos(new_doc):
    """
    Agrega una tabla de contenidos al documento proporcionado en el estilo 'Normal'.
    
    Args:
    new_doc (Document): El documento al que se añadirá la tabla de contenidos.
    """
    # Crear un párrafo para el título de la tabla de contenidos y centrarlo
    para = new_doc.add_paragraph("Tabla de Contenidos")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.name = 'Century Gothic'
        run.font.size = Pt(11)
        run.bold = True
        run.underline = True
        run.font.color.rgb = RGBColor(0, 32, 96)

    # Crear un párrafo vacío para insertar el campo de la tabla de contenidos
    paragraph = new_doc.add_paragraph()
    paragraph.paragraph_format.space_before = Inches(0)
    paragraph.paragraph_format.space_after = Inches(0)
    run = paragraph.add_run()
    
    # Insertar el campo de la tabla de contenidos
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
 
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need
 
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
 
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Haga clic derecho para actualizar el campo."
    fldChar3 = OxmlElement('w:updateFields')
    fldChar3.set(qn('w:val'), 'true')
    fldChar2.append(fldChar3)
 
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
 
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
 
    p_element = paragraph._p

# Función para agregar encabezados y pies de página
def add_header_footer(doc: Document, header_image_left: str, header_image_right: str, footer_image: str, footer_text: str):
    """
    Agrega un encabezado con dos imágenes (una a la izquierda y otra a la derecha) y un pie de página con una imagen a la derecha, texto a la izquierda y el número de página en el centro.
    
    Args:
    doc (Document): El documento al que se añadirán el encabezado y el pie de página.
    header_image_left (str): Ruta de la imagen para el encabezado (izquierda).
    header_image_right (str): Ruta de la imagen para el encabezado (derecha).
    footer_image (str): Ruta de la imagen para el pie de página.
    footer_text (str): Texto para el pie de página.
    """
    section = doc.sections[0]

    # Encabezado
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=doc.sections[0].page_width)
    header_table.autofit = True

    # Imagen izquierda en el encabezado
    header_cell_left = header_table.cell(0, 0)
    header_paragraph_left = header_cell_left.paragraphs[0]
    header_run_left = header_paragraph_left.add_run()
    header_run_left.add_picture(header_image_left, width=Inches(2.0))

    # Imagen derecha en el encabezado
    header_cell_right = header_table.cell(0, 1)
    header_paragraph_right = header_cell_right.paragraphs[0]
    header_paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header_run_right = header_paragraph_right.add_run()
    header_run_right.add_picture(header_image_right, width=Inches(2.0))

    # Pie de página
    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=3, width=doc.sections[0].page_width)
    footer_table.autofit = True

    # Texto izquierda en el pie de página
    footer_cell_left = footer_table.cell(0, 0)
    footer_paragraph_left = footer_cell_left.paragraphs[0]
    footer_paragraph_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    footer_run_left = footer_paragraph_left.add_run(footer_text)
    footer_run_left.font.size = Pt(8)

    # Número de página en el centro del pie de página
    footer_cell_center = footer_table.cell(0, 1)
    footer_paragraph_center = footer_cell_center.paragraphs[0]
    footer_paragraph_center.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run_center = footer_paragraph_center.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    footer_run_center._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    footer_run_center._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    footer_run_center._r.append(fldChar2)
    # Establecer el tamaño de la fuente del número de página
    footer_run_center.font.size = Pt(8)

    # Imagen derecha en el pie de página
    footer_cell_right = footer_table.cell(0, 2)
    footer_paragraph_right = footer_cell_right.paragraphs[0]
    footer_paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    footer_run_right = footer_paragraph_right.add_run()
    footer_run_right.add_picture(footer_image, width=Inches(2.0))


#############
# CONTINENTES
#############
def create_document_continentes(df_total, df_nme, file_path, titulo, fecha, header_image_left, header_image_right, footer_image, year_cerrado, year_corrido):
    """
    Crea un documento Word con los resultados proporcionados.

    Args:
    resultados (dict): Diccionario con los resultados que se incluirán en el informe.
    file_path (str): Ruta donde se guardará el documento de salida.
    titulo (str): El título principal del informe.
    fecha (str): La fecha de actualización del informe.
    header_image_left (str): Ruta de la imagen para el encabezado (izquierda).
    header_image_right (str): Ruta de la imagen para el encabezado (derecha).
    footer_image (str): Ruta de la imagen para el pie de página.
    footer_text (str): Texto para el pie de página.
    year_cerrado (str): Texto con el año cerrado del informe.
    year_corrido (str): Texto con el año corrido del informe.
    """
    doc = Document()
    estilos(doc)

    # Agregar encabezado y pie de página
    footer_text = """Calle 28 # 13ª - 15, Edificio CCI Pisos 35 - 36 | Bogotá, Colombia T: +57 (1) 560 0100 | info@procolombia.co | www.procolombia.co"""

    add_header_footer(doc, header_image_left, header_image_right, footer_image, footer_text)
        
    # Agregar el título principal del informe
    title_paragraph = doc.add_paragraph(f'TRES EJES CONTINENTES: {str(titulo).upper()}', style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar la fecha
    date_paragraph = doc.add_paragraph(f'ÚLTIMA ACTUALIZACIÓN: {fecha.upper()}', style='Title')
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    
    # Agregar la tabla de contenido
    doc.add_page_break()
    agregar_tabla_contenidos(doc)   
    doc.add_page_break()
        

    ###############
    # Exportaciones
    ###############
    add_heading(doc, 'Exportaciones', level=2, style='Heading 1')
    
    #####################
    # Tipo de exportación
    #####################
    add_heading(doc, 'Tipo de exportación', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_total["Resumen USD Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_total["Resumen USD Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    #####################################
    # Exportaciones no minero-energéticas
    #####################################
    add_heading(doc, 'Exportaciones no minero-energéticas', level=2, style='Heading 1')

    ###########
    #  Destinos
    ###########
    add_heading(doc, 'Destinos', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][2][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][2][1]), 'Table Grid')
    doc.add_paragraph()
    
    ########################
    # Departamento de origen
    ########################
    add_heading(doc, 'Departamento de origen', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][3][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][3][1]), 'Table Grid')
    doc.add_paragraph()

    ########
    # Sector
    ########
    add_heading(doc, 'Sector', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    #############
    # SubSectores
    #############
    add_heading(doc, 'Subsector', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Tablas Subsectores Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Tablas Subsectores Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    ####################
    # Conteo de Empresas
    ####################
    add_heading(doc, 'Empresas', level=3, style='Heading 2')
    # Año cerrado
    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
    run_cerrado = paragraph_cerrado.add_run(f'{str(df_nme["Numero de Empresas Cerrado"][0][1])} empresas')
    run_cerrado.bold = True
    # Año corrido
    paragraph_corrdo = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
    run_corrido = paragraph_corrdo.add_run(f'{str(df_nme["Numero de Empresas Corrido"][0][1])} empresas')
    run_corrido.bold = True

    ################
    # Datos empresas
    ################
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Tablas Empresas Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Tablas Empresas Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    # Guardar el documento
    doc.save(file_path)


######
# HUBS
######
def create_document_hub(df_total, df_nme, file_path, titulo, fecha, header_image_left, header_image_right, footer_image, year_cerrado, year_corrido):
    """
    Crea un documento Word con los resultados proporcionados.

    Args:
    resultados (dict): Diccionario con los resultados que se incluirán en el informe.
    file_path (str): Ruta donde se guardará el documento de salida.
    titulo (str): El título principal del informe.
    fecha (str): La fecha de actualización del informe.
    header_image_left (str): Ruta de la imagen para el encabezado (izquierda).
    header_image_right (str): Ruta de la imagen para el encabezado (derecha).
    footer_image (str): Ruta de la imagen para el pie de página.
    footer_text (str): Texto para el pie de página.
    year_cerrado (str): Texto con el año cerrado del informe.
    year_corrido (str): Texto con el año corrido del informe.
    """
    doc = Document()
    estilos(doc)

    # Agregar encabezado y pie de página
    footer_text = """Calle 28 # 13ª - 15, Edificio CCI Pisos 35 - 36 | Bogotá, Colombia T: +57 (1) 560 0100 | info@procolombia.co | www.procolombia.co"""

    add_header_footer(doc, header_image_left, header_image_right, footer_image, footer_text)
        
    # Agregar el título principal del informe
    title_paragraph = doc.add_paragraph(f'TRES EJES HUBs: {str(titulo).upper()}', style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar la fecha
    date_paragraph = doc.add_paragraph(f'ÚLTIMA ACTUALIZACIÓN: {fecha.upper()}', style='Title')
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Agregar la tabla de contenido
    doc.add_page_break()
    agregar_tabla_contenidos(doc)   
    doc.add_page_break()

    ###############
    # Exportaciones
    ###############
    add_heading(doc, 'Exportaciones', level=2, style='Heading 1')
    
    #####################
    # Tipo de exportación
    #####################
    add_heading(doc, 'Tipo de exportación', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_total["Resumen USD Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_total["Resumen USD Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    #####################################
    # Exportaciones no minero-energéticas
    #####################################
    add_heading(doc, 'Exportaciones no minero-energéticas', level=2, style='Heading 1')

    ###########
    #  Destinos
    ###########
    add_heading(doc, 'Destinos', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][2][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][2][1]), 'Table Grid')
    doc.add_paragraph()
    
    ########################
    # Departamento de origen
    ########################
    add_heading(doc, 'Departamento de origen', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][3][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][3][1]), 'Table Grid')
    doc.add_paragraph()

    ########
    # Sector
    ########
    add_heading(doc, 'Sector', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    #############
    # SubSectores
    #############
    add_heading(doc, 'Subsector', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Tablas Subsectores Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Tablas Subsectores Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    ####################
    # Conteo de Empresas
    ####################
    add_heading(doc, 'Empresas', level=3, style='Heading 2')
    # Año cerrado
    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
    run_cerrado = paragraph_cerrado.add_run(f'{str(df_nme["Numero de Empresas Cerrado"][0][1])} empresas')
    run_cerrado.bold = True
    # Año corrido
    paragraph_corrdo = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
    run_corrido = paragraph_corrdo.add_run(f'{str(df_nme["Numero de Empresas Corrido"][0][1])} empresas')
    run_corrido.bold = True

    ################
    # Datos empresas
    ################
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Tablas Empresas Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Tablas Empresas Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    # Guardar el documento
    doc.save(file_path)


########
# PAÍSES
########
def create_document_pais(df_total, df_nme, file_path, titulo, fecha, header_image_left, header_image_right, footer_image, year_cerrado, year_corrido):
    """
    Crea un documento Word con los resultados proporcionados.

    Args:
    resultados (dict): Diccionario con los resultados que se incluirán en el informe.
    file_path (str): Ruta donde se guardará el documento de salida.
    titulo (str): El título principal del informe.
    fecha (str): La fecha de actualización del informe.
    header_image_left (str): Ruta de la imagen para el encabezado (izquierda).
    header_image_right (str): Ruta de la imagen para el encabezado (derecha).
    footer_image (str): Ruta de la imagen para el pie de página.
    footer_text (str): Texto para el pie de página.
    year_cerrado (str): Texto con el año cerrado del informe.
    year_corrido (str): Texto con el año corrido del informe.
    """
    doc = Document()
    estilos(doc)

    # Agregar encabezado y pie de página
    footer_text = """Calle 28 # 13ª - 15, Edificio CCI Pisos 35 - 36 | Bogotá, Colombia T: +57 (1) 560 0100 | info@procolombia.co | www.procolombia.co"""

    add_header_footer(doc, header_image_left, header_image_right, footer_image, footer_text)
        
    # Agregar el título principal del informe
    title_paragraph = doc.add_paragraph(f'TRES EJES PAÍSES: {str(titulo).upper()}', style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar la fecha
    date_paragraph = doc.add_paragraph(f'ÚLTIMA ACTUALIZACIÓN: {fecha.upper()}', style='Title')
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    # Agregar la tabla de contenido
    doc.add_page_break()
    agregar_tabla_contenidos(doc)   
    doc.add_page_break()

    ###############
    # Exportaciones
    ###############
    add_heading(doc, 'Exportaciones', level=2, style='Heading 1')
    
    #####################
    # Tipo de exportación
    #####################
    add_heading(doc, 'Tipo de exportación', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_total["Resumen USD Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_total["Resumen USD Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    #####################################
    # Exportaciones no minero-energéticas
    #####################################
    add_heading(doc, 'Exportaciones no minero-energéticas', level=2, style='Heading 1')
  
    ########################
    # Departamento de origen
    ########################
    add_heading(doc, 'Departamento de origen', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][3][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][3][1]), 'Table Grid')
    doc.add_paragraph()

    ########
    # Sector
    ########
    add_heading(doc, 'Sector', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    #############
    # SubSectores
    #############
    add_heading(doc, 'Subsector', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Tablas Subsectores Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Tablas Subsectores Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    ####################
    # Conteo de Empresas
    ####################
    add_heading(doc, 'Empresas', level=3, style='Heading 2')
    # Año cerrado
    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
    run_cerrado = paragraph_cerrado.add_run(f'{str(df_nme["Numero de Empresas Cerrado"][0][1])} empresas')
    run_cerrado.bold = True
    # Año corrido
    paragraph_corrdo = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
    run_corrido = paragraph_corrdo.add_run(f'{str(df_nme["Numero de Empresas Corrido"][0][1])} empresas')
    run_corrido.bold = True

    ################
    # Datos empresas
    ################
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Tablas Empresas Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Tablas Empresas Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    # Guardar el documento
    doc.save(file_path)


###############
# DEPARTAMENTOS
###############
def create_document_departamento(df_total, df_nme, file_path, titulo, fecha, header_image_left, header_image_right, footer_image, year_cerrado, year_corrido):
    """
    Crea un documento Word con los resultados proporcionados.

    Args:
    resultados (dict): Diccionario con los resultados que se incluirán en el informe.
    file_path (str): Ruta donde se guardará el documento de salida.
    titulo (str): El título principal del informe.
    fecha (str): La fecha de actualización del informe.
    header_image_left (str): Ruta de la imagen para el encabezado (izquierda).
    header_image_right (str): Ruta de la imagen para el encabezado (derecha).
    footer_image (str): Ruta de la imagen para el pie de página.
    footer_text (str): Texto para el pie de página.
    year_cerrado (str): Texto con el año cerrado del informe.
    year_corrido (str): Texto con el año corrido del informe.
    """
    doc = Document()
    estilos(doc)

    # Agregar encabezado y pie de página
    footer_text = """Calle 28 # 13ª - 15, Edificio CCI Pisos 35 - 36 | Bogotá, Colombia T: +57 (1) 560 0100 | info@procolombia.co | www.procolombia.co"""

    add_header_footer(doc, header_image_left, header_image_right, footer_image, footer_text)
        
    # Agregar el título principal del informe
    title_paragraph = doc.add_paragraph(f'TRES EJES DEPARTAMENTOS: {str(titulo).upper()}', style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar la fecha
    date_paragraph = doc.add_paragraph(f'ÚLTIMA ACTUALIZACIÓN: {fecha.upper()}', style='Title')
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    # Agregar la tabla de contenido
    doc.add_page_break()
    agregar_tabla_contenidos(doc)   
    doc.add_page_break()

    ###############
    # Exportaciones
    ###############
    add_heading(doc, 'Exportaciones', level=2, style='Heading 1')
    
    #####################
    # Tipo de exportación
    #####################
    add_heading(doc, 'Tipo de exportación', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_total["Resumen USD Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_total["Resumen USD Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    #####################################
    # Exportaciones no minero-energéticas
    #####################################
    add_heading(doc, 'Exportaciones no minero-energéticas', level=2, style='Heading 1')

    ###########
    #  Destinos
    ###########
    add_heading(doc, 'Destinos', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][2][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][2][1]), 'Table Grid')
    doc.add_paragraph()
    
    ########
    # Sector
    ########
    add_heading(doc, 'Sector', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    #############
    # SubSectores
    #############
    add_heading(doc, 'Subsector', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Tablas Subsectores Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Tablas Subsectores Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    ######
    # TLCS
    ######
    add_heading(doc, 'Tratados de libre comercio', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][6][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][6][1]), 'Table Grid')
    doc.add_paragraph()

    ####################
    # Conteo de Empresas
    ####################
    add_heading(doc, 'Empresas', level=3, style='Heading 2')
    # Año cerrado
    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
    run_cerrado = paragraph_cerrado.add_run(f'{str(df_nme["Numero de Empresas Cerrado"][0][1])} empresas')
    run_cerrado.bold = True
    # Año corrido
    paragraph_corrdo = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
    run_corrido = paragraph_corrdo.add_run(f'{str(df_nme["Numero de Empresas Corrido"][0][1])} empresas')
    run_corrido.bold = True

    ################
    # Datos empresas
    ################
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Tablas Empresas Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Tablas Empresas Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    # Guardar el documento
    doc.save(file_path)

##########
# COLOMBIA
##########
def create_document_colombia(df_total, df_nme, file_path, fecha, header_image_left, header_image_right, footer_image, year_cerrado, year_corrido):
    """
    Crea un documento Word con los resultados proporcionados.

    Args:
    resultados (dict): Diccionario con los resultados que se incluirán en el informe.
    file_path (str): Ruta donde se guardará el documento de salida.
    fecha (str): La fecha de actualización del informe.
    header_image_left (str): Ruta de la imagen para el encabezado (izquierda).
    header_image_right (str): Ruta de la imagen para el encabezado (derecha).
    footer_image (str): Ruta de la imagen para el pie de página.
    footer_text (str): Texto para el pie de página.
    year_cerrado (str): Texto con el año cerrado del informe.
    year_corrido (str): Texto con el año corrido del informe.
    """
    doc = Document()
    estilos(doc)

    # Agregar encabezado y pie de página
    footer_text = """Calle 28 # 13ª - 15, Edificio CCI Pisos 35 - 36 | Bogotá, Colombia T: +57 (1) 560 0100 | info@procolombia.co | www.procolombia.co"""

    add_header_footer(doc, header_image_left, header_image_right, footer_image, footer_text)
        
    # Agregar el título principal del informe
    title_paragraph = doc.add_paragraph(f'TRES EJES COLOMBIA', style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar la fecha
    date_paragraph = doc.add_paragraph(f'ÚLTIMA ACTUALIZACIÓN: {fecha.upper()}', style='Title')
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Agregar la tabla de contenido
    doc.add_page_break()
    agregar_tabla_contenidos(doc)   
    doc.add_page_break()

    ###############
    # Exportaciones
    ###############
    add_heading(doc, 'Exportaciones', level=2, style='Heading 1')
    
    #####################
    # Tipo de exportación
    #####################
    add_heading(doc, 'Tipo de exportación', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_total["Resumen USD Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_total["Resumen USD Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    #####################################
    # Exportaciones no minero-energéticas
    #####################################
    add_heading(doc, 'Exportaciones no minero-energéticas', level=2, style='Heading 1')

    ###########
    #  Destinos
    ###########
    add_heading(doc, 'Destinos', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][2][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][2][1]), 'Table Grid')
    doc.add_paragraph()
    
    ########################
    # Departamento de origen
    ########################
    add_heading(doc, 'Departamento de origen', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][3][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][3][1]), 'Table Grid')
    doc.add_paragraph()

    ########
    # Sector
    ########
    add_heading(doc, 'Sector', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Resumen USD Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    #############
    # SubSectores
    #############
    add_heading(doc, 'Subsector', level=3, style='Heading 2')
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Tablas Subsectores Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Tablas Subsectores Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    ####################
    # Conteo de Empresas
    ####################
    add_heading(doc, 'Empresas', level=3, style='Heading 2')
    # Año cerrado
    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
    run_cerrado = paragraph_cerrado.add_run(f'{str(df_nme["Numero de Empresas Cerrado"][0][1])} empresas')
    run_cerrado.bold = True
    # Año corrido
    paragraph_corrdo = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
    run_corrido = paragraph_corrdo.add_run(f'{str(df_nme["Numero de Empresas Corrido"][0][1])} empresas')
    run_corrido.bold = True

    ################
    # Datos empresas
    ################
    # Año cerrado
    add_table(doc, pd.DataFrame(df_nme["Tablas Empresas Cerrado"][0][1]), 'Table Grid')
    doc.add_paragraph()
    # Año corrido
    add_table(doc, pd.DataFrame(df_nme["Tablas Empresas Corrido"][0][1]), 'Table Grid')
    doc.add_paragraph()

    # Guardar el documento
    doc.save(file_path)